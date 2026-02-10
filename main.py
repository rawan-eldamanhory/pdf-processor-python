"""
PDF Processor — CLI & Demo Runner.
"""

import argparse
import os
import sys
from pathlib import Path
from typing import List

# ── Create sample PDFs for the demo ──────────────────────────────────────────
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.units import cm
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors


def _make_sample_pdfs(out_dir: str = "sample_pdfs") -> dict:
    """
    Create three sample PDFs for demo purposes.
    """
    os.makedirs(out_dir, exist_ok=True)
    styles = getSampleStyleSheet()
    paths = {}

    # ---- sample_1.pdf : text-heavy ----------------------------------------
    p1 = os.path.join(out_dir, "sample_1.pdf")
    doc = SimpleDocTemplate(p1, pagesize=A4)
    story = []
    story.append(Paragraph("Annual Technology Report 2024", styles["Title"]))
    story.append(Spacer(1, 0.5 * cm))
    story.append(Paragraph("Executive Summary", styles["Heading1"]))
    story.append(Paragraph(
        "This report summarises the key technology trends observed in 2024. "
        "Artificial intelligence, cloud computing, and cybersecurity continue "
        "to dominate investment priorities across all industry sectors. "
        "Organisations that adopted automation early reported 35% higher "
        "efficiency gains compared to late adopters. Python remains the most "
        "popular language for data science and AI workloads for the fifth "
        "consecutive year.", styles["Normal"]
    ))
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph("Key Findings", styles["Heading1"]))
    for finding in [
        "Cloud spending grew 28% year-over-year.",
        "78% of enterprises now use some form of AI/ML.",
        "Cybersecurity incidents increased by 14%.",
        "Remote work tooling investment stabilised after pandemic peaks.",
        "Open-source adoption reached an all-time high.",
    ]:
        story.append(Paragraph(f"• {finding}", styles["Normal"]))
    story.append(PageBreak())
    story.append(Paragraph("Detailed Analysis", styles["Heading1"]))
    story.append(Paragraph(
        "The following sections provide a deep-dive into each technology "
        "category. Data was collected from 1,200 organisations across "
        "40 countries between January and October 2024. All monetary "
        "figures are in USD unless otherwise stated.", styles["Normal"]
    ))
    doc.build(story)
    paths["sample_1"] = p1
    print(f"Created {p1}")

    # ---- sample_2.pdf : tables + data ----------------------------------------
    p2 = os.path.join(out_dir, "sample_2.pdf")
    doc2 = SimpleDocTemplate(p2, pagesize=A4)
    story2 = []
    story2.append(Paragraph("Q3 2024 Sales Data", styles["Title"]))
    story2.append(Spacer(1, 0.5 * cm))
    story2.append(Paragraph("Regional Performance", styles["Heading1"]))

    table_data = [
        ["Region", "Q1 ($k)", "Q2 ($k)", "Q3 ($k)", "Growth"],
        ["North America", "1,240", "1,380", "1,520", "+10.1%"],
        ["Europe", "890", "920", "1,050", "+14.1%"],
        ["Asia Pacific", "670", "740", "810", "+9.5%"],
        ["Latin America", "310", "340", "370", "+8.8%"],
        ["Middle East", "180", "195", "220", "+12.8%"],
        ["TOTAL", "3,290", "3,575", "3,970", "+11.1%"],
    ]

    tbl = Table(table_data, colWidths=[5 * cm, 3 * cm, 3 * cm, 3 * cm, 2.5 * cm])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1A237E")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#E3F2FD")),
        ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -2), [colors.white, colors.HexColor("#F5F5F5")]),
    ]))
    story2.append(tbl)
    story2.append(Spacer(1, 0.3 * cm))
    story2.append(Paragraph("Table 1: Quarterly revenue by region (USD thousands)",
                            styles["Normal"]))
    story2.append(Spacer(1, 0.5 * cm))
    story2.append(Paragraph("Product Category Breakdown", styles["Heading1"]))
    cat_data = [
        ["Product", "Units Sold", "Revenue ($k)", "Margin %"],
        ["Software Licences", "4,200", "2,100", "82%"],
        ["Professional Services", "—", "980", "45%"],
        ["Hardware", "1,850", "620", "22%"],
        ["Support Contracts", "3,100", "270", "91%"],
    ]
    tbl2 = Table(cat_data, colWidths=[5 * cm, 3.5 * cm, 3.5 * cm, 2.5 * cm])
    tbl2.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#3949AB")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F5F5F5")]),
    ]))
    story2.append(tbl2)
    doc2.build(story2)
    paths["sample_2"] = p2
    print(f"Created {p2}")

    # ---- sample_3.pdf : short doc for merge test ----------------------------
    p3 = os.path.join(out_dir, "sample_3.pdf")
    doc3 = SimpleDocTemplate(p3, pagesize=A4)
    story3 = [
        Paragraph("Appendix A: Glossary", styles["Title"]),
        Spacer(1, 0.5 * cm),
    ]
    terms = {
        "AI": "Artificial Intelligence",
        "ML": "Machine Learning",
        "API": "Application Programming Interface",
        "SaaS": "Software as a Service",
        "OCR": "Optical Character Recognition",
    }
    for term, definition in terms.items():
        story3.append(Paragraph(f"<b>{term}</b> — {definition}", styles["Normal"]))
        story3.append(Spacer(1, 0.15 * cm))
    doc3.build(story3)
    paths["sample_3"] = p3
    print(f"Created {p3}")

    return paths


# ── Demo sections ─────────────────────────────────────────────────────────────


def demo_text_extraction(pdf_paths: dict) -> None:
    """Demo: extract text from a PDF."""
    from pdf_processor import extract_text_all_pages, extract_text_by_page

    print("\n" + "=" * 60)
    print("  DEMO 1: TEXT EXTRACTION")
    print("=" * 60)

    pdf = pdf_paths["sample_1"]
    print(f"\nExtracting ALL text from: {pdf}")
    text = extract_text_all_pages(pdf)
    print(f"  Total characters extracted: {len(text):,}")
    print(f"\n  First 400 characters:\n  {'─'*50}")
    print("  " + text[:400].replace("\n", "\n  "))

    print(f"\nExtracting only PAGE 1 from: {pdf}")
    page1 = extract_text_by_page(pdf, 1)
    print(f"  Page 1 → {len(page1)} characters")
    print(f"  Preview: {page1[:200]!r}")

    # Save to file
    out = "outputs/extracted_text.txt"
    os.makedirs("outputs", exist_ok=True)
    with open(out, "w", encoding="utf-8") as f:
        f.write(text)
    print(f"\n Saved full text → {out}")


def demo_table_extraction(pdf_paths: dict) -> None:
    """Demo: extract tables from a PDF."""
    from pdf_processor import extract_tables, tables_to_text

    print("\n" + "=" * 60)
    print("  DEMO 2: TABLE EXTRACTION")
    print("=" * 60)

    pdf = pdf_paths["sample_2"]
    print(f"\nExtracting tables from: {pdf}")
    tables = extract_tables(pdf)
    print(f"  Found {len(tables)} table(s)")

    for t in tables:
        print(f"\n  Table {t['index']} on page {t['page']}")
        print(f"    Columns : {t['headers']}")
        print(f"    Rows    : {len(t['rows'])}")

    formatted = tables_to_text(tables)
    out = "outputs/extracted_tables.txt"
    with open(out, "w", encoding="utf-8") as f:
        f.write(formatted)
    print(f"\n  Saved formatted tables → {out}")
    print(f"\n  Preview:\n  {'─'*50}")
    print("  " + formatted[:600].replace("\n", "\n  "))


def demo_merge_split(pdf_paths: dict) -> None:
    """Demo: merge multiple PDFs and split a PDF into pages."""
    from pdf_processor import merge_pdfs, split_pdf, extract_page_range

    print("\n" + "=" * 60)
    print("  DEMO 3: MERGE & SPLIT")
    print("=" * 60)

    # Merge
    print("\n🔗 Merging 3 sample PDFs into one…")
    merged_path = "outputs/merged.pdf"
    total = merge_pdfs(
        [pdf_paths["sample_1"], pdf_paths["sample_2"], pdf_paths["sample_3"]],
        merged_path
    )
    size = os.path.getsize(merged_path) // 1024
    print(f"Merged → {merged_path}  ({total} pages, {size} KB)")

    # Split
    print("\nSplitting sample_1.pdf into individual pages…")
    pages = split_pdf(pdf_paths["sample_1"], "outputs/split_pages")
    print(f" Split into {len(pages)} file(s):")
    for p in pages:
        print(f"     {p}")

    # Extract range
    print("\nExtracting pages 1–1 from sample_2.pdf…")
    ranged = extract_page_range(pdf_paths["sample_2"], 1, 1, "outputs/sample2_p1.pdf")
    print(f" Saved → {ranged}")


def demo_metadata(pdf_paths: dict) -> None:
    """Demo: read PDF metadata."""
    from pdf_processor import get_pdf_info

    print("\n" + "=" * 60)
    print("  DEMO 4: PDF METADATA")
    print("=" * 60)

    for name, path in pdf_paths.items():
        info = get_pdf_info(path)
        print(f"\n  {name}.pdf")
        print(f"     Pages     : {info['num_pages']}")
        print(f"     Size      : {info['file_size_kb']} KB")
        print(f"     Encrypted : {info['encrypted']}")
        print(f"     Title     : {info['title'] or '(none)'}")
        print(f"     Author    : {info['author'] or '(none)'}")


def demo_search(pdf_paths: dict) -> None:
    """Demo: search for text inside a PDF."""
    from pdf_processor import search_pdf

    print("\n" + "=" * 60)
    print("  DEMO 5: SEARCH INSIDE PDF")
    print("=" * 60)

    queries = ["Python", "cloud", r"\d+%"]
    pdf = pdf_paths["sample_1"]

    for q in queries:
        hits = search_pdf(pdf, q, case_sensitive=False)
        print(f"\n  Query: '{q}'  → {len(hits)} match(es)")
        for hit in hits[:3]:
            print(f"     Page {hit['page']}, line {hit['line_number']}: …{hit['line'][:60]}…")


def demo_report_generation() -> None:
    """Demo: create professional PDF reports."""
    from report_generator import generate_full_report

    print("\n" + "=" * 60)
    print("  DEMO 6: PROFESSIONAL REPORT GENERATION")
    print("=" * 60)

    out = "outputs/professional_report.pdf"
    print(f"\n Generating full report → {out}")

    generate_full_report(
        title="Annual Technology Report 2024",
        subtitle="AI, Cloud & Cybersecurity Deep Dive",
        author="PDF Processor Demo",
        output_path=out,
        sections=[
            {
                "heading": "Executive Summary",
                "level": 1,
                "body": (
                    "This report summarises the major technology trends "
                    "across 1,200 organisations worldwide in 2024. "
                    "Artificial intelligence adoption reached an inflection "
                    "point, with 78% of enterprises now running at least one "
                    "production ML workload. Cloud spending grew 28% YoY."
                ),
            },
            {
                "heading": "Key Highlights",
                "level": 1,
                "bullets": [
                    "Cloud spending grew 28% year-over-year to $680 billion.",
                    "78% of enterprises use AI/ML in production workloads.",
                    "Cybersecurity incidents increased 14% globally.",
                    "Open-source adoption reached an all-time high.",
                    "Python retained its #1 data science language ranking.",
                ],
            },
            {
                "heading": "Cloud Computing",
                "level": 2,
                "body": (
                    "Hyperscaler revenues (AWS, Azure, GCP) combined reached "
                    "$380 billion, representing 56% of total cloud spend. "
                    "Infrastructure-as-Code tooling became standard practice "
                    "in 91% of surveyed DevOps teams."
                ),
            },
            {
                "heading": "Artificial Intelligence",
                "level": 2,
                "body": (
                    "Generative AI saw explosive growth with 340% more "
                    "enterprise pilot programmes in 2024 than 2023. "
                    "However, only 23% of pilots reached production, "
                    "highlighting the gap between experimentation and value "
                    "delivery."
                ),
            },
        ],
        tables=[
            {
                "heading": "Regional Sales Performance",
                "headers": ["Region", "Q1 ($k)", "Q2 ($k)", "Q3 ($k)", "Growth"],
                "rows": [
                    ["North America", "1,240", "1,380", "1,520", "+10.1%"],
                    ["Europe", "890", "920", "1,050", "+14.1%"],
                    ["Asia Pacific", "670", "740", "810", "+9.5%"],
                    ["Latin America", "310", "340", "370", "+8.8%"],
                    ["Middle East", "180", "195", "220", "+12.8%"],
                    ["TOTAL", "3,290", "3,575", "3,970", "+11.1%"],
                ],
                "caption": "Table 1 — Q1–Q3 2024 Revenue by Region (USD thousands)",
            },
            {
                "heading": "Product Category Performance",
                "headers": ["Product", "Units Sold", "Revenue ($k)", "Margin %"],
                "rows": [
                    ["Software Licences", "4,200", "2,100", "82%"],
                    ["Professional Services", "—", "980", "45%"],
                    ["Hardware", "1,850", "620", "22%"],
                    ["Support Contracts", "3,100", "270", "91%"],
                ],
                "caption": "Table 2 — Q3 2024 Revenue by Product Category",
            },
        ],
        summary=(
            "2024 was a landmark year for enterprise technology adoption. "
            "Organisations that invested in cloud-native architectures and "
            "AI tooling outperformed peers by an average of 18% on key "
            "efficiency metrics. The outlook for 2025 remains strong, with "
            "continued growth in AI infrastructure and edge computing."
        ),
    )

    size = os.path.getsize(out) // 1024
    print(f" Report generated  ({size} KB)")


def demo_encrypt_decrypt(pdf_paths: dict) -> None:
    """Demo: encrypt and decrypt a PDF."""
    from pdf_processor import encrypt_pdf, decrypt_pdf

    print("\n" + "=" * 60)
    print("  DEMO 7: ENCRYPT & DECRYPT")
    print("=" * 60)

    src = pdf_paths["sample_3"]
    enc = "outputs/encrypted.pdf"
    dec = "outputs/decrypted.pdf"

    print(f"\n Encrypting {src} with password 'secret123'…")
    encrypt_pdf(src, enc, user_password="secret123")
    print(f"  Encrypted → {enc}")

    print(f"\n Decrypting {enc}…")
    decrypt_pdf(enc, dec, password="secret123")
    print(f"  Decrypted → {dec}")


def demo_rotate_watermark(pdf_paths: dict) -> None:
    """Demo: rotate pages."""
    from pdf_processor import rotate_pages
    from report_generator import generate_text_report

    print("\n" + "=" * 60)
    print("  DEMO 8: ROTATE PAGES")
    print("=" * 60)

    src = pdf_paths["sample_1"]
    out = "outputs/rotated.pdf"
    print(f"\n Rotating page 1 of {src} by 90°…")
    rotate_pages(src, out, degrees=90, page_numbers=[1])
    print(f"  Rotated → {out}")


def demo_docx(sample_dir: str = "sample_pdfs") -> None:
    """Demo: read a Word document and convert it to PDF."""
    from docx import Document
    from pdf_processor import extract_text_from_docx

    print("\n" + "=" * 60)
    print("  DEMO 9: WORD (.docx) PROCESSING")
    print("=" * 60)

    # Create a sample .docx
    docx_path = os.path.join(sample_dir, "sample.docx")
    doc = Document()
    doc.add_heading("Sample Word Document", 0)
    doc.add_paragraph(
        "This is a Word document created by python-docx. "
        "The PDF Processor can read its content and convert it to PDF."
    )
    doc.add_heading("Features", level=1)
    for f in [
        "Extract all paragraph text",
        "Convert to PDF report",
        "Works alongside PDF operations",
    ]:
        doc.add_paragraph(f, style="List Bullet")
    doc.save(docx_path)
    print(f" Created sample Word doc → {docx_path}")

    text = extract_text_from_docx(docx_path)
    print(f"\n  Extracted text ({len(text)} chars):")
    print(f"  {text[:300]!r}")

    out_pdf = "outputs/from_docx.pdf"
    from pdf_processor import docx_to_pdf_text_report
    docx_to_pdf_text_report(docx_path, out_pdf)
    print(f"\n Converted .docx → {out_pdf}")


# ── Main entry point ──────────────────────────────────────────────────────────


def main() -> None:
    """Run the full demo or a specific feature."""
    parser = argparse.ArgumentParser(
        description="PDF/Document Processor Demo"
    )
    parser.add_argument(
        "--feature",
        choices=["text", "table", "merge", "metadata", "search",
                 "report", "encrypt", "rotate", "docx", "all"],
        default="all",
        help="Which demo to run (default: all)",
    )
    args = parser.parse_args()

    print()
    print("=" * 60)
    print("  PDF / DOCUMENT PROCESSOR")
    print("  Full Feature Demo")
    print("=" * 60)

    print("\n Creating sample PDFs…")
    pdf_paths = _make_sample_pdfs()

    os.makedirs("outputs", exist_ok=True)

    feature = args.feature

    if feature in ("text", "all"):
        demo_text_extraction(pdf_paths)

    if feature in ("table", "all"):
        demo_table_extraction(pdf_paths)

    if feature in ("merge", "all"):
        demo_merge_split(pdf_paths)

    if feature in ("metadata", "all"):
        demo_metadata(pdf_paths)

    if feature in ("search", "all"):
        demo_search(pdf_paths)

    if feature in ("report", "all"):
        demo_report_generation()

    if feature in ("encrypt", "all"):
        demo_encrypt_decrypt(pdf_paths)

    if feature in ("rotate", "all"):
        demo_rotate_watermark(pdf_paths)

    if feature in ("docx", "all"):
        demo_docx()

    # ── Final summary ─────────────────────────────────────────────────────────
    print("\n" + "=" * 60)
    print("  ALL DEMOS COMPLETE")
    print("=" * 60)

    print("\n Output files created:")
    for root, _dirs, files in os.walk("outputs"):
        for fname in sorted(files):
            fpath = os.path.join(root, fname)
            size = os.path.getsize(fpath) // 1024
            print(f"  {fpath:<45} {size:>5} KB")

    print("  • Text extraction with pdfplumber + pypdf fallback")
    print("  • Table extraction → structured data")
    print("  • Merge, split, and page-range extraction")
    print("  • PDF metadata access")
    print("  • Full-text search with regex")
    print("  • Professional report generation with ReportLab")
    print("  • Encryption / decryption")
    print("  • Page rotation")
    print("  • Word (.docx) reading and conversion")
    print()


if __name__ == "__main__":
    main()
