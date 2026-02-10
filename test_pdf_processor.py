"""
Test suite for PDF/Document Processor.
"""

import os
import tempfile
from pathlib import Path
from typing import Generator

import pytest
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from reportlab.lib import colors


# ─────────────────────────────────────────────────────────────────────────────
# Fixtures
# ─────────────────────────────────────────────────────────────────────────────


def _create_simple_pdf(path: str, text: str = "Hello PDF World") -> str:
    """Helper: create a minimal PDF at *path*."""
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(path, pagesize=A4)
    story = [Paragraph(text, styles["Normal"])]
    doc.build(story)
    return path


def _create_table_pdf(path: str) -> str:
    """Helper: create a PDF that contains a visible table."""
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(path, pagesize=A4)

    table_data = [
        ["Name", "Score", "Grade"],
        ["Alice", "95", "A"],
        ["Bob", "82", "B"],
        ["Carol", "78", "C"],
    ]
    tbl = Table(table_data)
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
    ]))

    doc.build([
        Paragraph("Test Table Document", styles["Title"]),
        Spacer(1, 20),
        tbl,
    ])
    return path


@pytest.fixture
def tmp_dir() -> Generator[str, None, None]:
    """Temporary directory cleaned up after each test."""
    with tempfile.TemporaryDirectory() as d:
        yield d


@pytest.fixture
def simple_pdf(tmp_dir: str) -> str:
    """A simple single-page PDF."""
    path = os.path.join(tmp_dir, "simple.pdf")
    return _create_simple_pdf(path, "Python PDF Processing is fun!")


@pytest.fixture
def multi_page_pdf(tmp_dir: str) -> str:
    """A multi-page PDF."""
    from reportlab.platypus import PageBreak

    path = os.path.join(tmp_dir, "multi.pdf")
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(path, pagesize=A4)
    story = []
    for i in range(1, 4):
        story.append(Paragraph(f"Page {i} Content", styles["Title"]))
        story.append(Paragraph(
            f"This is the text on page {i}. Python is great.", styles["Normal"]
        ))
        if i < 3:
            story.append(PageBreak())
    doc.build(story)
    return path


@pytest.fixture
def table_pdf(tmp_dir: str) -> str:
    """A PDF that contains a table."""
    path = os.path.join(tmp_dir, "table.pdf")
    return _create_table_pdf(path)


@pytest.fixture
def sample_docx(tmp_dir: str) -> str:
    """A sample Word document."""
    from docx import Document

    path = os.path.join(tmp_dir, "test.docx")
    doc = Document()
    doc.add_heading("Test Document", 0)
    doc.add_paragraph("First paragraph text.")
    doc.add_paragraph("Second paragraph text.")
    doc.save(path)
    return path


# ─────────────────────────────────────────────────────────────────────────────
# pdf_processor — text extraction
# ─────────────────────────────────────────────────────────────────────────────


class TestTextExtraction:
    """Tests for text extraction functions."""

    def test_extract_text_all_pages_returns_string(self, simple_pdf: str):
        from pdf_processor import extract_text_all_pages

        text = extract_text_all_pages(simple_pdf)
        assert isinstance(text, str)

    def test_extract_text_all_pages_contains_content(self, simple_pdf: str):
        from pdf_processor import extract_text_all_pages

        text = extract_text_all_pages(simple_pdf)
        assert "Python PDF Processing" in text

    def test_extract_text_all_pages_multipage(self, multi_page_pdf: str):
        from pdf_processor import extract_text_all_pages

        text = extract_text_all_pages(multi_page_pdf)
        assert "Page 1" in text
        assert "Page 2" in text
        assert "Page 3" in text

    def test_extract_text_all_pages_page_markers(self, multi_page_pdf: str):
        from pdf_processor import extract_text_all_pages

        text = extract_text_all_pages(multi_page_pdf)
        assert "--- Page 1 ---" in text

    def test_extract_text_file_not_found(self, tmp_dir: str):
        from pdf_processor import extract_text_all_pages

        with pytest.raises(FileNotFoundError):
            extract_text_all_pages(os.path.join(tmp_dir, "missing.pdf"))

    def test_extract_text_by_page_first_page(self, multi_page_pdf: str):
        from pdf_processor import extract_text_by_page

        text = extract_text_by_page(multi_page_pdf, 1)
        assert isinstance(text, str)

    def test_extract_text_by_page_valid_content(self, multi_page_pdf: str):
        from pdf_processor import extract_text_by_page

        text = extract_text_by_page(multi_page_pdf, 2)
        assert "Page 2" in text

    def test_extract_text_by_page_out_of_range(self, simple_pdf: str):
        from pdf_processor import extract_text_by_page

        with pytest.raises(ValueError):
            extract_text_by_page(simple_pdf, 999)

    def test_extract_text_by_page_zero_raises(self, simple_pdf: str):
        from pdf_processor import extract_text_by_page

        with pytest.raises(ValueError):
            extract_text_by_page(simple_pdf, 0)

    def test_extract_text_page_range_returns_dict(self, multi_page_pdf: str):
        from pdf_processor import extract_text_page_range

        result = extract_text_page_range(multi_page_pdf, 1, 2)
        assert isinstance(result, dict)
        assert 1 in result
        assert 2 in result

    def test_extract_text_page_range_clips_to_total(self, multi_page_pdf: str):
        from pdf_processor import extract_text_page_range

        result = extract_text_page_range(multi_page_pdf, 2, 999)
        assert len(result) == 2  # pages 2 and 3


# ─────────────────────────────────────────────────────────────────────────────
# pdf_processor — table extraction
# ─────────────────────────────────────────────────────────────────────────────


class TestTableExtraction:
    """Tests for table extraction."""

    def test_extract_tables_returns_list(self, table_pdf: str):
        from pdf_processor import extract_tables

        tables = extract_tables(table_pdf)
        assert isinstance(tables, list)

    def test_extract_tables_finds_at_least_one(self, table_pdf: str):
        from pdf_processor import extract_tables

        tables = extract_tables(table_pdf)
        assert len(tables) >= 1

    def test_extract_tables_has_headers(self, table_pdf: str):
        from pdf_processor import extract_tables

        tables = extract_tables(table_pdf)
        if tables:
            assert "headers" in tables[0]
            assert isinstance(tables[0]["headers"], list)

    def test_extract_tables_has_rows(self, table_pdf: str):
        from pdf_processor import extract_tables

        tables = extract_tables(table_pdf)
        if tables:
            assert "rows" in tables[0]
            assert isinstance(tables[0]["rows"], list)

    def test_extract_tables_page_key_present(self, table_pdf: str):
        from pdf_processor import extract_tables

        tables = extract_tables(table_pdf)
        if tables:
            assert "page" in tables[0]

    def test_tables_to_text_returns_string(self, table_pdf: str):
        from pdf_processor import extract_tables, tables_to_text

        tables = extract_tables(table_pdf)
        text = tables_to_text(tables)
        assert isinstance(text, str)

    def test_tables_to_text_empty_input(self):
        from pdf_processor import tables_to_text

        assert tables_to_text([]) == ""


# ─────────────────────────────────────────────────────────────────────────────
# pdf_processor — merge & split
# ─────────────────────────────────────────────────────────────────────────────


class TestMergeSplit:
    """Tests for merge and split operations."""

    def test_merge_pdfs_creates_file(self, simple_pdf: str, tmp_dir: str):
        from pdf_processor import merge_pdfs

        out = os.path.join(tmp_dir, "merged.pdf")
        merge_pdfs([simple_pdf, simple_pdf], out)
        assert os.path.exists(out)

    def test_merge_pdfs_returns_page_count(
        self, simple_pdf: str, multi_page_pdf: str, tmp_dir: str
    ):
        from pdf_processor import merge_pdfs

        out = os.path.join(tmp_dir, "merged.pdf")
        count = merge_pdfs([simple_pdf, multi_page_pdf], out)
        assert count >= 4  # 1 + 3

    def test_merge_pdfs_missing_file_raises(self, simple_pdf: str, tmp_dir: str):
        from pdf_processor import merge_pdfs

        out = os.path.join(tmp_dir, "merged.pdf")
        with pytest.raises(FileNotFoundError):
            merge_pdfs([simple_pdf, "/nonexistent/path.pdf"], out)

    def test_split_pdf_creates_files(self, multi_page_pdf: str, tmp_dir: str):
        from pdf_processor import split_pdf

        out_dir = os.path.join(tmp_dir, "split")
        files = split_pdf(multi_page_pdf, out_dir)
        assert len(files) == 3
        for f in files:
            assert os.path.exists(f)

    def test_split_pdf_chunks(self, multi_page_pdf: str, tmp_dir: str):
        from pdf_processor import split_pdf

        out_dir = os.path.join(tmp_dir, "chunks")
        files = split_pdf(multi_page_pdf, out_dir, pages_per_chunk=2)
        # 3 pages / 2 per chunk = 2 files
        assert len(files) == 2

    def test_extract_page_range_creates_file(
        self, multi_page_pdf: str, tmp_dir: str
    ):
        from pdf_processor import extract_page_range

        out = os.path.join(tmp_dir, "range.pdf")
        result = extract_page_range(multi_page_pdf, 1, 2, out)
        assert os.path.exists(result)


# ─────────────────────────────────────────────────────────────────────────────
# pdf_processor — metadata, rotate, encrypt
# ─────────────────────────────────────────────────────────────────────────────


class TestMetadataAndSecurity:
    """Tests for metadata, rotate, and encryption."""

    def test_get_pdf_info_num_pages(self, multi_page_pdf: str):
        from pdf_processor import get_pdf_info

        info = get_pdf_info(multi_page_pdf)
        assert info["num_pages"] == 3

    def test_get_pdf_info_keys(self, simple_pdf: str):
        from pdf_processor import get_pdf_info

        info = get_pdf_info(simple_pdf)
        for key in ("num_pages", "file_size_kb", "encrypted", "title", "author"):
            assert key in info

    def test_get_pdf_info_not_encrypted_by_default(self, simple_pdf: str):
        from pdf_processor import get_pdf_info

        assert not get_pdf_info(simple_pdf)["encrypted"]

    def test_get_pdf_info_file_size_positive(self, simple_pdf: str):
        from pdf_processor import get_pdf_info

        assert get_pdf_info(simple_pdf)["file_size_kb"] > 0

    def test_rotate_pages_creates_file(self, simple_pdf: str, tmp_dir: str):
        from pdf_processor import rotate_pages

        out = os.path.join(tmp_dir, "rotated.pdf")
        result = rotate_pages(simple_pdf, out, degrees=90)
        assert os.path.exists(result)

    def test_encrypt_decrypt_roundtrip(self, simple_pdf: str, tmp_dir: str):
        from pdf_processor import decrypt_pdf, encrypt_pdf

        enc = os.path.join(tmp_dir, "enc.pdf")
        dec = os.path.join(tmp_dir, "dec.pdf")

        encrypt_pdf(simple_pdf, enc, user_password="test123")
        assert os.path.exists(enc)

        decrypt_pdf(enc, dec, password="test123")
        assert os.path.exists(dec)

    def test_decrypt_wrong_password_raises(self, simple_pdf: str, tmp_dir: str):
        from pdf_processor import decrypt_pdf, encrypt_pdf

        enc = os.path.join(tmp_dir, "enc2.pdf")
        encrypt_pdf(simple_pdf, enc, user_password="correct")

        with pytest.raises(ValueError):
            decrypt_pdf(enc, os.path.join(tmp_dir, "dec2.pdf"), password="wrong")


# ─────────────────────────────────────────────────────────────────────────────
# pdf_processor — search
# ─────────────────────────────────────────────────────────────────────────────


class TestSearch:
    """Tests for PDF search."""

    def test_search_finds_match(self, simple_pdf: str):
        from pdf_processor import search_pdf

        hits = search_pdf(simple_pdf, "Python")
        assert len(hits) >= 1

    def test_search_returns_page_key(self, simple_pdf: str):
        from pdf_processor import search_pdf

        hits = search_pdf(simple_pdf, "PDF")
        if hits:
            assert "page" in hits[0]
            assert "line" in hits[0]
            assert "match" in hits[0]

    def test_search_case_insensitive(self, simple_pdf: str):
        from pdf_processor import search_pdf

        hits_lower = search_pdf(simple_pdf, "python", case_sensitive=False)
        hits_upper = search_pdf(simple_pdf, "PYTHON", case_sensitive=False)
        assert len(hits_lower) == len(hits_upper)

    def test_search_no_match_returns_empty(self, simple_pdf: str):
        from pdf_processor import search_pdf

        hits = search_pdf(simple_pdf, "ZZZNOMATCH999")
        assert hits == []

    def test_search_regex_pattern(self, simple_pdf: str):
        from pdf_processor import search_pdf

        # Regex: any word ending in 'ing'
        hits = search_pdf(simple_pdf, r"\w+ing")
        assert isinstance(hits, list)


# ─────────────────────────────────────────────────────────────────────────────
# pdf_processor — Word (.docx)
# ─────────────────────────────────────────────────────────────────────────────


class TestDocx:
    """Tests for Word document processing."""

    def test_extract_text_from_docx_returns_string(self, sample_docx: str):
        from pdf_processor import extract_text_from_docx

        text = extract_text_from_docx(sample_docx)
        assert isinstance(text, str)

    def test_extract_text_from_docx_contains_content(self, sample_docx: str):
        from pdf_processor import extract_text_from_docx

        text = extract_text_from_docx(sample_docx)
        assert "Test Document" in text or "First paragraph" in text

    def test_docx_to_pdf_creates_file(self, sample_docx: str, tmp_dir: str):
        from pdf_processor import docx_to_pdf_text_report

        out = os.path.join(tmp_dir, "converted.pdf")
        result = docx_to_pdf_text_report(sample_docx, out)
        assert os.path.exists(result)


# ─────────────────────────────────────────────────────────────────────────────
# report_generator
# ─────────────────────────────────────────────────────────────────────────────


class TestReportGenerator:
    """Tests for PDF report generation."""

    def test_generate_text_report_creates_file(self, tmp_dir: str):
        from report_generator import generate_text_report

        out = os.path.join(tmp_dir, "text_report.pdf")
        result = generate_text_report(
            title="Test Report",
            sections=[{"heading": "Section 1", "body": "Some body text."}],
            output_path=out,
        )
        assert os.path.exists(result)

    def test_generate_text_report_file_not_empty(self, tmp_dir: str):
        from report_generator import generate_text_report

        out = os.path.join(tmp_dir, "text_report2.pdf")
        generate_text_report(
            title="Test",
            sections=[{"heading": "H", "body": "B"}],
            output_path=out,
        )
        assert os.path.getsize(out) > 1000

    def test_generate_text_report_with_bullets(self, tmp_dir: str):
        from report_generator import generate_text_report

        out = os.path.join(tmp_dir, "bullets.pdf")
        result = generate_text_report(
            title="Bullet Report",
            sections=[
                {
                    "heading": "List",
                    "bullets": ["Item A", "Item B", "Item C"],
                }
            ],
            output_path=out,
        )
        assert os.path.exists(result)

    def test_generate_table_report_creates_file(self, tmp_dir: str):
        from report_generator import generate_table_report

        out = os.path.join(tmp_dir, "table_report.pdf")
        result = generate_table_report(
            title="Table Report",
            tables=[
                {
                    "heading": "Sales",
                    "headers": ["Month", "Revenue"],
                    "rows": [["Jan", "$100k"], ["Feb", "$120k"]],
                }
            ],
            output_path=out,
        )
        assert os.path.exists(result)

    def test_generate_full_report_creates_file(self, tmp_dir: str):
        from report_generator import generate_full_report

        out = os.path.join(tmp_dir, "full_report.pdf")
        result = generate_full_report(
            title="Full Report",
            output_path=out,
            sections=[{"heading": "Intro", "body": "Introduction text."}],
            tables=[
                {
                    "heading": "Data",
                    "headers": ["A", "B"],
                    "rows": [["1", "2"]],
                }
            ],
            summary="This is the summary.",
        )
        assert os.path.exists(result)

    def test_generate_full_report_with_author(self, tmp_dir: str):
        from report_generator import generate_full_report

        out = os.path.join(tmp_dir, "authored.pdf")
        result = generate_full_report(
            title="Authored Report",
            output_path=out,
            author="Test Author",
            subtitle="A Test Subtitle",
        )
        assert os.path.exists(result)

    def test_generate_full_report_file_size_reasonable(self, tmp_dir: str):
        from report_generator import generate_full_report

        out = os.path.join(tmp_dir, "size_check.pdf")
        generate_full_report(
            title="Size Test",
            output_path=out,
            sections=[{"heading": "Section", "body": "Body " * 100}],
        )
        size_kb = os.path.getsize(out) // 1024
        assert 1 <= size_kb <= 500  # sane range

    @pytest.mark.parametrize("n_rows", [0, 1, 5, 20])
    def test_generate_table_report_varying_rows(
        self, tmp_dir: str, n_rows: int
    ):
        from report_generator import generate_table_report

        out = os.path.join(tmp_dir, f"table_{n_rows}.pdf")
        rows = [[str(i), f"value_{i}"] for i in range(n_rows)]
        if n_rows == 0:
            rows = []
        generate_table_report(
            title=f"Table {n_rows} rows",
            tables=[
                {
                    "heading": "Test",
                    "headers": ["ID", "Value"],
                    "rows": rows,
                }
            ],
            output_path=out,
        )
        assert os.path.exists(out)


# ─────────────────────────────────────────────────────────────────────────────
# Smoke tests — integration
# ─────────────────────────────────────────────────────────────────────────────


class TestSmoke:
    """End-to-end smoke tests."""

    @pytest.mark.smoke
    def test_full_pipeline(self, tmp_dir: str):
        """Create PDF → extract text → generate report."""
        from pdf_processor import extract_text_all_pages
        from report_generator import generate_text_report

        src = os.path.join(tmp_dir, "source.pdf")
        _create_simple_pdf(src, "Smoke test source document with Python content.")

        text = extract_text_all_pages(src)
        assert len(text) > 0

        out = os.path.join(tmp_dir, "smoke_report.pdf")
        generate_text_report(
            title="Smoke Report",
            sections=[{"heading": "Extracted", "body": text}],
            output_path=out,
        )
        assert os.path.exists(out)

    @pytest.mark.smoke
    def test_merge_split_roundtrip(self, tmp_dir: str):
        """Merge three PDFs, then split back, check file count."""
        from pdf_processor import merge_pdfs, split_pdf

        a = _create_simple_pdf(os.path.join(tmp_dir, "a.pdf"), "A page")
        b = _create_simple_pdf(os.path.join(tmp_dir, "b.pdf"), "B page")
        c = _create_simple_pdf(os.path.join(tmp_dir, "c.pdf"), "C page")

        merged = os.path.join(tmp_dir, "merged.pdf")
        merge_pdfs([a, b, c], merged)

        pages = split_pdf(merged, os.path.join(tmp_dir, "pages"))
        assert len(pages) == 3
