"""
Standalone test runner for PDF Processor.
"""

import os
import sys
import tempfile
import traceback
import unittest

# ─────────────────────────────────────────────────────────────────────────────
# Helpers shared by tests
# ─────────────────────────────────────────────────────────────────────────────


def _make_pdf(path: str, text: str = "Hello World Python Test") -> str:
    """Create a minimal PDF at *path* using ReportLab."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate

    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(path, pagesize=A4)
    doc.build([Paragraph(text, styles["Normal"])])
    return path


def _make_multi_pdf(path: str) -> str:
    """Create a 3-page PDF."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate

    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(path, pagesize=A4)
    story = []
    for i in range(1, 4):
        story.append(Paragraph(f"Page {i} Python content here", styles["Normal"]))
        if i < 3:
            story.append(PageBreak())
    doc.build(story)
    return path


def _make_table_pdf(path: str) -> str:
    """Create a PDF with a table."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(path, pagesize=A4)
    data = [
        ["Name", "Score", "Grade"],
        ["Alice", "95", "A"],
        ["Bob", "82", "B"],
    ]
    tbl = Table(data)
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
    ]))
    doc.build([Paragraph("Table Document", styles["Title"]), Spacer(1, 20), tbl])
    return path


# ─────────────────────────────────────────────────────────────────────────────
# Test classes
# ─────────────────────────────────────────────────────────────────────────────


class TestTextExtraction(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.mkdtemp()
        self.pdf = _make_pdf(os.path.join(self.tmp, "simple.pdf"))
        self.multi = _make_multi_pdf(os.path.join(self.tmp, "multi.pdf"))

    def test_extract_all_returns_string(self):
        from pdf_processor import extract_text_all_pages
        text = extract_text_all_pages(self.pdf)
        self.assertIsInstance(text, str)

    def test_extract_all_contains_content(self):
        from pdf_processor import extract_text_all_pages
        text = extract_text_all_pages(self.pdf)
        self.assertIn("Python", text)

    def test_extract_all_multipage_has_markers(self):
        from pdf_processor import extract_text_all_pages
        text = extract_text_all_pages(self.multi)
        self.assertIn("--- Page 1 ---", text)

    def test_extract_all_file_not_found_raises(self):
        from pdf_processor import extract_text_all_pages
        with self.assertRaises(FileNotFoundError):
            extract_text_all_pages("/no/such/file.pdf")

    def test_extract_by_page_returns_string(self):
        from pdf_processor import extract_text_by_page
        text = extract_text_by_page(self.multi, 1)
        self.assertIsInstance(text, str)

    def test_extract_by_page_correct_content(self):
        from pdf_processor import extract_text_by_page
        text = extract_text_by_page(self.multi, 2)
        self.assertIn("Page 2", text)

    def test_extract_by_page_out_of_range_raises(self):
        from pdf_processor import extract_text_by_page
        with self.assertRaises(ValueError):
            extract_text_by_page(self.pdf, 999)

    def test_extract_by_page_zero_raises(self):
        from pdf_processor import extract_text_by_page
        with self.assertRaises(ValueError):
            extract_text_by_page(self.pdf, 0)

    def test_extract_page_range_returns_dict(self):
        from pdf_processor import extract_text_page_range
        result = extract_text_page_range(self.multi, 1, 2)
        self.assertIsInstance(result, dict)
        self.assertIn(1, result)
        self.assertIn(2, result)

    def test_extract_page_range_clips(self):
        from pdf_processor import extract_text_page_range
        result = extract_text_page_range(self.multi, 2, 999)
        self.assertEqual(len(result), 2)


class TestTableExtraction(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.mkdtemp()
        self.pdf = _make_table_pdf(os.path.join(self.tmp, "table.pdf"))

    def test_extract_returns_list(self):
        from pdf_processor import extract_tables
        self.assertIsInstance(extract_tables(self.pdf), list)

    def test_extract_finds_table(self):
        from pdf_processor import extract_tables
        tables = extract_tables(self.pdf)
        self.assertGreaterEqual(len(tables), 1)

    def test_table_has_headers(self):
        from pdf_processor import extract_tables
        tables = extract_tables(self.pdf)
        if tables:
            self.assertIn("headers", tables[0])

    def test_table_has_rows(self):
        from pdf_processor import extract_tables
        tables = extract_tables(self.pdf)
        if tables:
            self.assertIn("rows", tables[0])

    def test_tables_to_text_returns_string(self):
        from pdf_processor import extract_tables, tables_to_text
        tables = extract_tables(self.pdf)
        self.assertIsInstance(tables_to_text(tables), str)

    def test_tables_to_text_empty(self):
        from pdf_processor import tables_to_text
        self.assertEqual(tables_to_text([]), "")


class TestMergeSplit(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.mkdtemp()
        self.a = _make_pdf(os.path.join(self.tmp, "a.pdf"), "File A")
        self.b = _make_pdf(os.path.join(self.tmp, "b.pdf"), "File B")
        self.multi = _make_multi_pdf(os.path.join(self.tmp, "multi.pdf"))

    def test_merge_creates_file(self):
        from pdf_processor import merge_pdfs
        out = os.path.join(self.tmp, "merged.pdf")
        merge_pdfs([self.a, self.b], out)
        self.assertTrue(os.path.exists(out))

    def test_merge_page_count(self):
        from pdf_processor import merge_pdfs
        out = os.path.join(self.tmp, "m.pdf")
        n = merge_pdfs([self.a, self.b], out)
        self.assertEqual(n, 2)

    def test_merge_missing_file_raises(self):
        from pdf_processor import merge_pdfs
        with self.assertRaises(FileNotFoundError):
            merge_pdfs([self.a, "/no/such/file.pdf"], os.path.join(self.tmp, "x.pdf"))

    def test_split_one_per_page(self):
        from pdf_processor import split_pdf
        files = split_pdf(self.multi, os.path.join(self.tmp, "split"))
        self.assertEqual(len(files), 3)
        for f in files:
            self.assertTrue(os.path.exists(f))

    def test_split_two_per_chunk(self):
        from pdf_processor import split_pdf
        files = split_pdf(self.multi, os.path.join(self.tmp, "chunks"), pages_per_chunk=2)
        self.assertEqual(len(files), 2)

    def test_extract_page_range_creates_file(self):
        from pdf_processor import extract_page_range
        out = os.path.join(self.tmp, "range.pdf")
        result = extract_page_range(self.multi, 1, 2, out)
        self.assertTrue(os.path.exists(result))


class TestMetadataAndSecurity(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.mkdtemp()
        self.pdf = _make_pdf(os.path.join(self.tmp, "base.pdf"))
        self.multi = _make_multi_pdf(os.path.join(self.tmp, "multi.pdf"))

    def test_get_info_page_count(self):
        from pdf_processor import get_pdf_info
        self.assertEqual(get_pdf_info(self.multi)["num_pages"], 3)

    def test_get_info_keys(self):
        from pdf_processor import get_pdf_info
        info = get_pdf_info(self.pdf)
        for k in ("num_pages", "file_size_kb", "encrypted", "title", "author"):
            self.assertIn(k, info)

    def test_get_info_not_encrypted(self):
        from pdf_processor import get_pdf_info
        self.assertFalse(get_pdf_info(self.pdf)["encrypted"])

    def test_get_info_positive_size(self):
        from pdf_processor import get_pdf_info
        self.assertGreater(get_pdf_info(self.pdf)["file_size_kb"], 0)

    def test_rotate_creates_file(self):
        from pdf_processor import rotate_pages
        out = os.path.join(self.tmp, "rotated.pdf")
        result = rotate_pages(self.pdf, out, degrees=90)
        self.assertTrue(os.path.exists(result))

    def test_encrypt_decrypt_roundtrip(self):
        from pdf_processor import decrypt_pdf, encrypt_pdf
        enc = os.path.join(self.tmp, "enc.pdf")
        dec = os.path.join(self.tmp, "dec.pdf")
        encrypt_pdf(self.pdf, enc, user_password="pass123")
        self.assertTrue(os.path.exists(enc))
        decrypt_pdf(enc, dec, password="pass123")
        self.assertTrue(os.path.exists(dec))

    def test_decrypt_wrong_password_raises(self):
        from pdf_processor import decrypt_pdf, encrypt_pdf
        enc = os.path.join(self.tmp, "enc2.pdf")
        encrypt_pdf(self.pdf, enc, user_password="correct")
        with self.assertRaises(ValueError):
            decrypt_pdf(enc, os.path.join(self.tmp, "dec2.pdf"), password="wrong")


class TestSearch(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.mkdtemp()
        self.pdf = _make_pdf(
            os.path.join(self.tmp, "search.pdf"),
            "Python is a programming language. Python is great!"
        )

    def test_finds_match(self):
        from pdf_processor import search_pdf
        hits = search_pdf(self.pdf, "Python")
        self.assertGreaterEqual(len(hits), 1)

    def test_hit_has_keys(self):
        from pdf_processor import search_pdf
        hits = search_pdf(self.pdf, "Python")
        self.assertIn("page", hits[0])
        self.assertIn("line", hits[0])
        self.assertIn("match", hits[0])

    def test_case_insensitive(self):
        from pdf_processor import search_pdf
        lower = search_pdf(self.pdf, "python", case_sensitive=False)
        upper = search_pdf(self.pdf, "PYTHON", case_sensitive=False)
        self.assertEqual(len(lower), len(upper))

    def test_no_match_empty_list(self):
        from pdf_processor import search_pdf
        self.assertEqual(search_pdf(self.pdf, "ZZZNOMATCH"), [])

    def test_regex_pattern(self):
        from pdf_processor import search_pdf
        hits = search_pdf(self.pdf, r"Py\w+")
        self.assertGreaterEqual(len(hits), 1)


class TestDocx(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.mkdtemp()
        from docx import Document
        path = os.path.join(self.tmp, "test.docx")
        doc = Document()
        doc.add_heading("Test Document", 0)
        doc.add_paragraph("First paragraph text in docx.")
        doc.save(path)
        self.docx_path = path

    def test_extract_returns_string(self):
        from pdf_processor import extract_text_from_docx
        self.assertIsInstance(extract_text_from_docx(self.docx_path), str)

    def test_extract_contains_content(self):
        from pdf_processor import extract_text_from_docx
        text = extract_text_from_docx(self.docx_path)
        self.assertTrue(len(text) > 0)

    def test_docx_to_pdf_creates_file(self):
        from pdf_processor import docx_to_pdf_text_report
        out = os.path.join(self.tmp, "converted.pdf")
        result = docx_to_pdf_text_report(self.docx_path, out)
        self.assertTrue(os.path.exists(result))


class TestReportGenerator(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.mkdtemp()

    def test_text_report_creates_file(self):
        from report_generator import generate_text_report
        out = os.path.join(self.tmp, "text.pdf")
        result = generate_text_report(
            title="Test",
            sections=[{"heading": "H1", "body": "Body text here."}],
            output_path=out
        )
        self.assertTrue(os.path.exists(result))

    def test_text_report_not_empty(self):
        from report_generator import generate_text_report
        out = os.path.join(self.tmp, "text2.pdf")
        generate_text_report("T", [{"body": "X"}], out)
        self.assertGreater(os.path.getsize(out), 1000)

    def test_text_report_with_bullets(self):
        from report_generator import generate_text_report
        out = os.path.join(self.tmp, "bullets.pdf")
        result = generate_text_report(
            title="Bullets",
            sections=[{"heading": "List", "bullets": ["A", "B", "C"]}],
            output_path=out
        )
        self.assertTrue(os.path.exists(result))

    def test_table_report_creates_file(self):
        from report_generator import generate_table_report
        out = os.path.join(self.tmp, "table.pdf")
        result = generate_table_report(
            title="Tables",
            tables=[{
                "heading": "Sales",
                "headers": ["Month", "Revenue"],
                "rows": [["Jan", "$100k"], ["Feb", "$120k"]],
            }],
            output_path=out
        )
        self.assertTrue(os.path.exists(result))

    def test_full_report_creates_file(self):
        from report_generator import generate_full_report
        out = os.path.join(self.tmp, "full.pdf")
        result = generate_full_report(
            title="Full",
            output_path=out,
            author="Tester",
            subtitle="Subtitle here",
            sections=[{"heading": "Intro", "body": "Intro text."}],
            tables=[{
                "heading": "Data",
                "headers": ["A", "B"],
                "rows": [["1", "2"]],
            }],
            summary="Summary text."
        )
        self.assertTrue(os.path.exists(result))

    def test_full_report_size_reasonable(self):
        from report_generator import generate_full_report
        out = os.path.join(self.tmp, "size.pdf")
        generate_full_report(
            title="Size",
            output_path=out,
            sections=[{"heading": "S", "body": "B " * 100}]
        )
        kb = os.path.getsize(out) // 1024
        self.assertGreaterEqual(kb, 1)
        self.assertLessEqual(kb, 500)


class TestSmoke(unittest.TestCase):
    """End-to-end smoke tests."""

    def test_full_pipeline(self):
        with tempfile.TemporaryDirectory() as d:
            from pdf_processor import extract_text_all_pages
            from report_generator import generate_text_report

            src = _make_pdf(os.path.join(d, "src.pdf"), "Smoke test Python content.")
            text = extract_text_all_pages(src)
            self.assertGreater(len(text), 0)

            out = os.path.join(d, "report.pdf")
            generate_text_report("Smoke", [{"heading": "X", "body": text}], out)
            self.assertTrue(os.path.exists(out))

    def test_merge_split_roundtrip(self):
        with tempfile.TemporaryDirectory() as d:
            from pdf_processor import merge_pdfs, split_pdf

            a = _make_pdf(os.path.join(d, "a.pdf"), "A")
            b = _make_pdf(os.path.join(d, "b.pdf"), "B")
            c = _make_pdf(os.path.join(d, "c.pdf"), "C")

            merged = os.path.join(d, "merged.pdf")
            merge_pdfs([a, b, c], merged)

            pages = split_pdf(merged, os.path.join(d, "pages"))
            self.assertEqual(len(pages), 3)


# ─────────────────────────────────────────────────────────────────────────────
# Runner
# ─────────────────────────────────────────────────────────────────────────────


def run_tests():
    print()
    print("=" * 70)
    print("  PDF PROCESSOR — TEST SUITE")
    print("=" * 70)
    print()

    loader = unittest.TestLoader()
    suite = unittest.TestSuite()

    test_classes = [
        TestTextExtraction,
        TestTableExtraction,
        TestMergeSplit,
        TestMetadataAndSecurity,
        TestSearch,
        TestDocx,
        TestReportGenerator,
        TestSmoke,
    ]

    for cls in test_classes:
        suite.addTests(loader.loadTestsFromTestCase(cls))

    runner = unittest.TextTestRunner(verbosity=2, stream=sys.stdout)
    result = runner.run(suite)

    print()
    print("=" * 70)
    total = result.testsRun
    passed = total - len(result.failures) - len(result.errors)
    print(f"  RESULTS: {passed}/{total} tests passed")

    if result.failures:
        print(f"\n  FAILURES ({len(result.failures)}):")
        for test, tb in result.failures:
            print(f"     {test}")

    if result.errors:
        print(f"\n  ERRORS ({len(result.errors)}):")
        for test, tb in result.errors:
            print(f"     {test}")

    if result.wasSuccessful():
        print("\n ALL TESTS PASSED")
    else:
        print("\n Some tests failed. Please review above.")

    print("=" * 70)
    return 0 if result.wasSuccessful() else 1


if __name__ == "__main__":
    sys.exit(run_tests())
